"""
Layer 0: Pre-flight, Parsing & Pre-computation

- Format routing (.pdf -> PyMuPDF, .docx -> python-docx)
- Scanned/image-based PDF detection
- Raw text + spatial block extraction
- Section detection
- Role & bullet parsing
- Date extraction + interval merging -> total_experience_years
- Career stage classification
- is_current_role derivation
- Bullet dataclass population with y_position and word_count
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from .constants import PRESENT_STRINGS, SECTION_KEYWORDS
from .models import (
    Bullet,
    CareerStage,
    ContactInfo,
    FontInfo,
    ParsedDocument,
    Role,
    Section,
    SpatialBlock,
)

# ---------------------------------------------------------------------------
# 0A  Format routing & scanned PDF detection
# ---------------------------------------------------------------------------

def parse_file(file_path: str) -> ParsedDocument:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(file_path)
    if ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    raise ValueError(f"Unsupported format: {ext}. Accepted: PDF, DOCX")


def _parse_pdf(path: str) -> ParsedDocument:
    doc = fitz.open(path)
    if not doc.page_count:
        raise ValueError("PDF has no pages.")

    first_page = doc[0]
    text_sample = first_page.get_text().strip()
    images = first_page.get_images()
    if len(text_sample) < 100 and len(images) > 0:
        raise ValueError(
            "PDF appears to be image-based (scanned). "
            "Cannot parse. Please provide a text-based PDF."
        )

    all_text_parts: list[str] = []
    spatial_blocks: list[SpatialBlock] = []
    page_height = first_page.rect.height
    page_width = first_page.rect.width

    for page_idx, page in enumerate(doc):
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    txt = span.get("text", "").strip()
                    if not txt:
                        continue
                    bbox = span.get("bbox", block.get("bbox", (0, 0, 0, 0)))
                    font_info = FontInfo(
                        size=span.get("size", 12),
                        name=span.get("font", ""),
                        is_bold="bold" in span.get("font", "").lower()
                                or (span.get("flags", 0) & 2 ** 4) != 0,
                        color=span.get("color", 0),
                    )
                    spatial_blocks.append(SpatialBlock(
                        text=txt,
                        bbox=tuple(bbox),
                        font=font_info,
                        page_num=page_idx,
                    ))
        all_text_parts.append(page.get_text())

    raw_text = "\n".join(all_text_parts)
    parsed = _build_document(raw_text, spatial_blocks)
    parsed.page_count = doc.page_count
    parsed.page_height = page_height
    parsed.page_width = page_width
    parsed.file_type = "pdf"
    doc.close()
    return parsed


def _parse_docx(path: str) -> ParsedDocument:
    document = DocxDocument(path)
    paragraphs_text: list[str] = []
    spatial_blocks: list[SpatialBlock] = []
    y_cursor = 0.0

    for para in document.paragraphs:
        txt = para.text.strip()
        if not txt:
            y_cursor += 14
            continue
        font_size = 12.0
        is_bold = False
        if para.runs:
            run = para.runs[0]
            if run.font.size:
                font_size = run.font.size.pt
            is_bold = bool(run.bold)
        spatial_blocks.append(SpatialBlock(
            text=txt,
            bbox=(72, y_cursor, 540, y_cursor + font_size + 4),
            font=FontInfo(size=font_size, name="", is_bold=is_bold),
            page_num=0,
        ))
        paragraphs_text.append(txt)
        y_cursor += font_size + 8

    raw_text = "\n".join(paragraphs_text)
    parsed = _build_document(raw_text, spatial_blocks)
    parsed.page_height = max(y_cursor, 842.0)
    parsed.file_type = "docx"

    # Check for hidden text boxes in DOCX XML
    parsed._docx_body = document.element.body  # stash for Layer 1
    return parsed


def parse_text_input(text: str) -> ParsedDocument:
    """For pasted JD or CV text (no spatial data)."""
    blocks: list[SpatialBlock] = []
    y = 0.0
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            y += 14
            continue
        blocks.append(SpatialBlock(
            text=line,
            bbox=(72, y, 540, y + 16),
            font=FontInfo(size=12, name=""),
            page_num=0,
        ))
        y += 16
    parsed = _build_document(text, blocks)
    parsed.file_type = "text"
    return parsed


# ---------------------------------------------------------------------------
# Build the ParsedDocument from raw text + spatial blocks
# ---------------------------------------------------------------------------

def _build_document(raw_text: str, spatial_blocks: list[SpatialBlock]) -> ParsedDocument:
    doc = ParsedDocument(raw_text=raw_text, spatial_blocks=spatial_blocks)
    doc.word_count = len(raw_text.split())

    # Detect sections
    doc.sections = _detect_sections(raw_text, spatial_blocks)

    # Extract contact info
    doc.contact = _extract_contact(raw_text)

    # Detect summary
    for sec in doc.sections:
        if sec.name == "summary":
            doc.has_summary = True
            doc.summary_text = sec.raw_text
            break

    # Extract roles and bullets
    doc.roles = _extract_roles(doc.sections, spatial_blocks)
    doc.bullets = []
    for role in doc.roles:
        doc.bullets.extend(role.bullets)

    # Pre-compute experience
    job_intervals = _extract_job_date_intervals(doc.roles)
    doc.total_experience_years = _calculate_experience_years(job_intervals)
    doc.career_stage = _classify_career_stage(doc.total_experience_years)

    return doc


# ---------------------------------------------------------------------------
# Section detection
# ---------------------------------------------------------------------------

def _detect_sections(raw_text: str, blocks: list[SpatialBlock]) -> list[Section]:
    lines = raw_text.split("\n")
    sections: list[Section] = []
    current_name = "header"
    current_lines: list[str] = []
    current_y = 0.0

    for line in lines:
        clean = line.strip()
        if not clean:
            continue
        lower = clean.lower().rstrip(":")
        detected = None
        for section_key, keywords in SECTION_KEYWORDS.items():
            if any(lower == kw or lower.startswith(kw) for kw in keywords):
                if len(clean) < 45:
                    detected = section_key
                    break

        if detected:
            if current_lines:
                y_pos = _find_block_y(current_name, blocks) or current_y
                sections.append(Section(
                    name=current_name,
                    raw_text="\n".join(current_lines).strip(),
                    y_position=y_pos,
                ))
            current_name = detected
            current_lines = []
            current_y = _find_block_y(clean, blocks) or current_y
        else:
            current_lines.append(clean)

    if current_lines:
        y_pos = _find_block_y(current_name, blocks) or current_y
        sections.append(Section(
            name=current_name,
            raw_text="\n".join(current_lines).strip(),
            y_position=y_pos,
        ))

    return sections


def _find_block_y(text_fragment: str, blocks: list[SpatialBlock]) -> Optional[float]:
    text_lower = text_fragment.lower().strip()
    for b in blocks:
        if text_lower in b.text.lower():
            return b.bbox[1]
    return None


# ---------------------------------------------------------------------------
# Contact extraction
# ---------------------------------------------------------------------------

def _extract_contact(text: str) -> ContactInfo:
    contact = ContactInfo()

    email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.]+', text)
    if email_match:
        contact.email = email_match.group()
        domain = contact.email.split("@")[1].split(".")[0].lower()
        from .constants import UNPROFESSIONAL_EMAIL_DOMAINS
        contact.email_professional = domain not in UNPROFESSIONAL_EMAIL_DOMAINS

    phone_match = re.search(r'[\+]?[\d\s\-\(\)]{7,15}', text)
    if phone_match:
        digits = re.sub(r'\D', '', phone_match.group())
        if len(digits) >= 7:
            contact.phone = phone_match.group().strip()

    linkedin_match = re.search(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+', text, re.I)
    if linkedin_match:
        contact.linkedin = linkedin_match.group()
        contact.has_linkedin = True

    location_match = re.search(
        r'(?:^|\n)\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
        text
    )
    if location_match:
        contact.location = location_match.group(1).strip()

    return contact


# ---------------------------------------------------------------------------
# Role & Bullet extraction
# ---------------------------------------------------------------------------

_DATE_PATTERN = re.compile(
    r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}'
    r'|(?:\d{1,2}[/\-]\d{4})'
    r'|\d{4})',
    re.I,
)

_ROLE_LINE_PATTERN = re.compile(
    r'^(.+?)\s*[\|\u2014\u2013\-,]+\s*(.+?)\s*[\|\u2014\u2013\-,]+\s*'
    r'((?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{4}'
    r'|\d{4})\s*[\u2013\u2014\-–]+\s*'
    r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{4}'
    r'|\d{4}|[Pp]resent|[Cc]urrent|[Nn]ow))',
    re.I | re.M,
)

_BULLET_CHARS = re.compile(r'^[\s]*[•\-\*►▪◦·‣⁃]\s*')
_NUMBERED_BULLET = re.compile(r'^[\s]*\d+[\.\)]\s*')


def _extract_roles(sections: list[Section], blocks: list[SpatialBlock]) -> list[Role]:
    roles: list[Role] = []
    exp_sections = [s for s in sections if s.name in ("experience", "projects")]
    if not exp_sections:
        return _fallback_role_extraction(sections, blocks)

    for section in exp_sections:
        section_roles = _parse_roles_from_text(section.raw_text, blocks)
        section.roles = section_roles
        roles.extend(section_roles)

    return roles


def _fallback_role_extraction(sections: list[Section], blocks: list[SpatialBlock]) -> list[Role]:
    """If no experience section found, try to find roles anywhere."""
    all_text = "\n".join(s.raw_text for s in sections)
    return _parse_roles_from_text(all_text, blocks)


def _parse_roles_from_text(text: str, blocks: list[SpatialBlock]) -> list[Role]:
    lines = text.split("\n")
    roles: list[Role] = []
    current_role: Optional[Role] = None
    bullet_buffer: list[str] = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        date_matches = _DATE_PATTERN.findall(line)
        has_dates = len(date_matches) >= 1

        is_bullet = bool(_BULLET_CHARS.match(line)) or bool(_NUMBERED_BULLET.match(line))
        clean_line = _BULLET_CHARS.sub("", _NUMBERED_BULLET.sub("", line)).strip()

        if has_dates and not is_bullet and len(line) < 120:
            if current_role and bullet_buffer:
                current_role.bullets = _make_bullets(
                    bullet_buffer, current_role, blocks
                )
            start_dt, end_dt, is_current = _parse_date_range(line)
            title, company = _extract_title_company(line, date_matches)
            current_role = Role(
                title=title,
                company=company,
                start_date=start_dt,
                end_date=end_dt,
                is_current=is_current,
            )
            if start_dt and end_dt:
                current_role.duration_months = (end_dt - start_dt).days / 30.44
            roles.append(current_role)
            bullet_buffer = []
        elif is_bullet or (current_role and len(clean_line) > 15):
            if clean_line:
                bullet_buffer.append(clean_line)

    if current_role and bullet_buffer:
        current_role.bullets = _make_bullets(bullet_buffer, current_role, blocks)

    return roles


def _make_bullets(
    texts: list[str], role: Role, blocks: list[SpatialBlock]
) -> list[Bullet]:
    bullets = []
    for text in texts:
        y_pos = 0.0
        for b in blocks:
            if text[:30].lower() in b.text.lower():
                y_pos = b.bbox[1]
                break
        bullets.append(Bullet(
            text=text,
            role=f"{role.company} — {role.title}",
            is_current_role=role.is_current,
            y_position=y_pos,
        ))
    return bullets


def _extract_title_company(line: str, date_matches: list[str]) -> tuple[str, str]:
    clean = line
    for dm in date_matches:
        clean = clean.replace(dm, "")
    clean = re.sub(r'[\u2013\u2014\-–]+\s*$', '', clean)
    clean = re.sub(r'[\u2013\u2014\-–]+\s*(?:present|current|now)\s*$', '', clean, flags=re.I)
    clean = clean.strip(" \t|,\u2013\u2014-–")

    parts = re.split(r'\s*[\|\u2014\u2013\-–,]+\s*', clean)
    parts = [p.strip() for p in parts if p.strip()]

    if len(parts) >= 2:
        return parts[0], parts[1]
    if parts:
        return parts[0], ""
    return "Unknown Role", ""


# ---------------------------------------------------------------------------
# Date parsing & interval merging
# ---------------------------------------------------------------------------

_MONTH_MAP = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}


def _parse_single_date(s: str) -> Optional[datetime]:
    s = s.strip()
    if not s:
        return None
    if s.lower() in PRESENT_STRINGS:
        return datetime.now()

    m = re.match(r'(\w+)\s+(\d{4})', s)
    if m:
        month_str = m.group(1)[:3].lower()
        year = int(m.group(2))
        month = _MONTH_MAP.get(month_str, 1)
        return datetime(year, month, 1)

    m = re.match(r'(\d{1,2})[/\-](\d{4})', s)
    if m:
        return datetime(int(m.group(2)), int(m.group(1)), 1)

    m = re.match(r'(\d{4})', s)
    if m:
        return datetime(int(m.group(1)), 1, 1)

    return None


def _parse_date_range(line: str) -> tuple[Optional[datetime], Optional[datetime], bool]:
    date_segments = re.findall(
        r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{4}'
        r'|\d{1,2}[/\-]\d{4}'
        r'|\d{4}'
        r'|[Pp]resent|[Cc]urrent|[Nn]ow',
        line, re.I,
    )
    if len(date_segments) >= 2:
        start = _parse_single_date(date_segments[-2])
        end_str = date_segments[-1]
        is_current = end_str.lower().strip() in PRESENT_STRINGS
        end = _parse_single_date(end_str)
        if is_current and end is None:
            end = datetime.now()
        return start, end, is_current
    if len(date_segments) == 1:
        start = _parse_single_date(date_segments[0])
        return start, start, False
    return None, None, False


def _extract_job_date_intervals(
    roles: list[Role],
) -> list[tuple[datetime, datetime]]:
    intervals = []
    for role in roles:
        if role.start_date and role.end_date:
            intervals.append((role.start_date, role.end_date))
    return intervals


def _calculate_experience_years(
    intervals: list[tuple[datetime, datetime]],
) -> float:
    if not intervals:
        return 0.0
    sorted_intervals = sorted(intervals)
    merged = [sorted_intervals[0]]
    for start, end in sorted_intervals[1:]:
        if start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    return sum((end - start).days for start, end in merged) / 365.25


def _classify_career_stage(years: float) -> CareerStage:
    if years < 2:
        return CareerStage.STUDENT
    if years <= 8:
        return CareerStage.MID
    return CareerStage.SENIOR
